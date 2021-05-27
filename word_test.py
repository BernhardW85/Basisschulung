from docx import Document

document = Document()

document.add_heading('Das ist mein Titel', 0)

p = document.add_paragraph('Test Absatz ')
p.add_run('bold Text').bold = True
p.add_run(' und der Text ist wieder nicht fett ')
p.add_run('und noch einen kursiven.').italic = True

document.add_heading('Überschrift 1', level=1)
document.add_heading('Überschrift 2', level=2)
document.add_paragraph('test quote', style='Intense Quote')

document.add_page_break()
document.add_paragraph('Text auf der neuen Seite', style='List Bullet')

document

document.save('test.docx')
